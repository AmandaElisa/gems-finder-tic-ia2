"""Converte os documentos Markdown de docs/ para .docx.

Existe porque colar Markdown no Google Docs não preserva formatação: os `#`
viram texto literal e as tabelas viram linhas de pipes. Um `.docx` aberto via
*Arquivo → Abrir → Upload* chega com títulos, negrito e tabelas de verdade.

Cobre o subconjunto de Markdown que os relatórios do projeto usam — títulos,
negrito, itálico, `código`, listas, tabelas e blocos de código. Não é um
conversor completo, e não precisa ser: um pandoc faria mais, mas exigiria
instalar um binário externo só para gerar dois arquivos.

    python scripts/md_para_docx.py docs/apresentacao/relatorio-metodologia.md
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# Trechos inline reconhecidos, na ordem em que precisam ser testados:
# o negrito vem antes do itálico porque `**` contém `*`.
INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def _escrever_inline(paragrafo, texto: str) -> None:
    """Escreve `texto` no parágrafo, aplicando negrito, itálico e código."""
    for parte in INLINE.split(texto):
        if not parte:
            continue
        if parte.startswith("**") and parte.endswith("**"):
            paragrafo.add_run(parte[2:-2]).bold = True
        elif parte.startswith("`") and parte.endswith("`"):
            run = paragrafo.add_run(parte[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        elif parte.startswith("*") and parte.endswith("*"):
            paragrafo.add_run(parte[1:-1]).italic = True
        else:
            paragrafo.add_run(parte)


def _linha_de_tabela(linha: str) -> list[str]:
    """Células de uma linha de tabela Markdown, sem os pipes das pontas."""
    return [c.strip() for c in linha.strip().strip("|").split("|")]


def _tabela(doc: Document, linhas: list[str]) -> None:
    """Uma tabela Markdown como tabela do Word, com a primeira linha em negrito."""
    corpo = [ln for ln in linhas if not re.fullmatch(r"\|[\s|:-]+\|", ln.strip())]
    if not corpo:
        return
    celulas = [_linha_de_tabela(ln) for ln in corpo]
    colunas = max(len(c) for c in celulas)
    tabela = doc.add_table(rows=0, cols=colunas)
    tabela.style = "Light Grid Accent 1"
    for i, linha in enumerate(celulas):
        celulas_docx = tabela.add_row().cells
        for j, conteudo in enumerate(linha[:colunas]):
            p = celulas_docx[j].paragraphs[0]
            _escrever_inline(p, conteudo)
            if i == 0:
                for run in p.runs:
                    run.bold = True


def converter(origem: Path, destino: Path) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    linhas = origem.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(linhas):
        linha = linhas[i]

        # Bloco de código: copiado literalmente até a cerca de fechamento.
        if linha.startswith("```"):
            i += 1
            codigo: list[str] = []
            while i < len(linhas) and not linhas[i].startswith("```"):
                codigo.append(linhas[i])
                i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(codigo))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            i += 1
            continue

        # Tabela: acumula as linhas contíguas que começam com pipe.
        if linha.strip().startswith("|"):
            bloco = []
            while i < len(linhas) and linhas[i].strip().startswith("|"):
                bloco.append(linhas[i])
                i += 1
            _tabela(doc, bloco)
            doc.add_paragraph()
            continue

        if linha.startswith("---") and set(linha.strip()) == {"-"}:
            doc.add_page_break()
        elif titulo := re.match(r"^(#{1,4})\s+(.*)", linha):
            nivel = len(titulo.group(1))
            _escrever_inline(doc.add_heading("", level=nivel), titulo.group(2))
        elif item := re.match(r"^\s*[-*]\s+(.*)", linha):
            _escrever_inline(doc.add_paragraph("", style="List Bullet"),
                             item.group(1))
        elif numerado := re.match(r"^\s*\d+\.\s+(.*)", linha):
            _escrever_inline(doc.add_paragraph("", style="List Number"),
                             numerado.group(1))
        elif citacao := re.match(r"^>\s*(.*)", linha):
            p = doc.add_paragraph("", style="Intense Quote")
            _escrever_inline(p, citacao.group(1))
        elif linha.strip():
            # Parágrafo: junta as linhas seguintes até a próxima linha vazia,
            # porque o Markdown quebra linha por largura e o Word não deve.
            texto = [linha.strip()]
            while (i + 1 < len(linhas) and linhas[i + 1].strip()
                   and not re.match(r"^(#{1,4}\s|\s*[-*]\s|\s*\d+\.\s|>|\||```)",
                                    linhas[i + 1])):
                i += 1
                texto.append(linhas[i].strip())
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _escrever_inline(p, " ".join(texto))
        i += 1

    doc.save(destino)
    print(f"{origem.name} -> {destino}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        sys.exit(__doc__)
    for caminho in sys.argv[1:]:
        origem = Path(caminho)
        converter(origem, origem.with_suffix(".docx"))
